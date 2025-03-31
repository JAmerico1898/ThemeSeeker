import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import googleapiclient.discovery
import googleapiclient.errors
import google.generativeai as genai
import os
from googleapiclient.discovery import build
import re
from bs4 import BeautifulSoup
import html
import io
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import requests
from io import BytesIO
import hashlib
import base64
import json
import zipfile

# Set page config
st.set_page_config(
    page_title="Spirituality Trends Analyzer",
    page_icon="✨",
    layout="wide"
)

# App title and description
st.title("Spirituality YouTube Trend Analyzer")
st.markdown("""
This application mines popular YouTube videos related to spirituality from different time periods 
and uses Gemini AI to suggest thematic content for lectures targeting various age groups.
""")

# Initialize session state for philosophy context and themes
if 'philosophy_context' not in st.session_state:
    st.session_state['philosophy_context'] = ""
if 'generated_themes' not in st.session_state:
    st.session_state['generated_themes'] = []
if 'selected_theme_index' not in st.session_state:
    st.session_state['selected_theme_index'] = None

# Sidebar for API keys
with st.sidebar:
    # Get API keys from secrets or environment variables
    try:
        youtube_api_key = st.secrets["youtube_api_key"]
        gemini_api_key = st.secrets["GOOGLE_API_KEY"]
    except:
        # Fallback to environment variables if secrets not available
        youtube_api_key = os.environ.get("YOUTUBE_API_KEY", "")
        gemini_api_key = os.environ.get("GEMINI_API_KEY", "")
    
    st.header("Search Parameters")
    search_query = st.text_input("Search Query", value="spirituality philosophy meaning of life")
    max_results = st.slider("Maximum Videos per Time Period", 5, 50, 20)
    
    # Instead of file upload, we'll load the HTML content from the provided file
    st.header("Philosophy Context")
    st.info("Philosophy context has been loaded from the Rosacruz Áurea website")
    
    # Process the HTML content that was provided
    try:
        # The HTML content is the one from the document uploaded previously
        content = """
<!DOCTYPE html>
<html class="html" lang="pt-BR">
<head>
    <!-- Head content omitted for brevity -->
</head>
<body class="home page-template-default page page-id-7 wp-embed-responsive oceanwp-theme dropdown-mobile default-breakpoint content-full-screen has-topbar page-header-disabled has-breadcrumbs elementor-default elementor-kit-4 elementor-page elementor-page-7">
    <!-- Body content from the Rosacruz Áurea website -->
    <div class="entry clr" itemprop="text">
        <h3>Rosacruz Áurea | LECTORIUM ROSICRUCIANUM</h3>
        <p>A Rosacruz Áurea é uma Escola iniciática contemporânea, dedicada à transformação da consciência e da vida do ser humano atual.</p>
        <p>A fonte do conhecimento da Rosacruz Áurea é a própria Sabedoria Universal, manifestada em todos os tempos, culturas e povos.</p>
        <p>A Rosacruz Áurea dirige-se ao ser humano buscador, oferecendo-lhe elementos para que ele encontre em si mesmo suas respostas e as converta em seu próprio caminho de transformação. Estes elementos também se encontram em seu símbolo: ponto central, triângulo, quadrado e círculo. Juntos, eles representam em todos os níveis, macrocósmico, cósmico ou microcósmico, um símbolo universal da criação divina.</p>
        <p>A consciência humana é prisioneira de seu próprio egocentrismo. Esse estado de consciência nunca será a base do processo de iniciação. Aqueles que desejam seguir o caminho da iniciação devem superar a si mesmos, pois aqueles que se superam tornam-se capazes de verdadeiramente amar e servir a humanidade e o mundo.</p>
        <p>O Caminho da Iniciação tem diferentes aspectos: 1. Autoconhecimento: tornar-se consciente do seu próprio egocentrismo; 2. Conexão: estabelecer a conexão inicial e consciente com o Ser Real e superar o egocentrismo; 3. Nova Consciência: através dessa conexão inicial, transformar o pensamento, o sentimento e a ação, permitindo o surgimento de uma nova consciência;</p>
        <p>4. Consciência Espiritual: por meio de um trabalho contínuo, a nova consciência se desenvolve e amadurece, unindo-se plenamente ao Ser Real; 5. Transmutação: a consciência espiritual desencadeia uma transformação energética dos veículos da personalidade, focando em seus aspectos mais sutis; 6. Nova Vitalidade: a transformação energética caminha para uma transformação orgânica, resultando em uma nova e total energia vital;</p>
        <p>7. Reintegração: reintegrado ao universo, o novo ser humano torna-se um servo do mundo e da humanidade. </p>
    </div>
</body>
</html>
        """
        
        # Store the raw HTML content
        st.session_state['philosophy_context'] = content
        
        # Extract text from HTML using BeautifulSoup
        soup = BeautifulSoup(content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        
        # Get text
        cleaned_text = soup.get_text()
        
        # Clean up the text
        lines = (line.strip() for line in cleaned_text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        cleaned_text = '\n'.join(chunk for chunk in chunks if chunk)
        
        # Store cleaned text
        st.session_state['philosophy_context_cleaned'] = cleaned_text
        
        # Show a sample of the extracted text
        with st.expander("Preview extracted text"):
            st.write(cleaned_text[:2700] + "..." if len(cleaned_text) > 2700 else cleaned_text)
            
    except Exception as e:
        st.error(f"Error processing philosophy context: {str(e)}")
        
    st.caption("Note: This application requires API keys to function properly.")

# Function to get date in ISO format for a given period
def get_date_for_period(period):
    today = datetime.now()
    
    if period == "1 week":
        past_date = today - timedelta(weeks=1)
    elif period == "1 month":
        past_date = today - timedelta(days=30)
    elif period == "6 months":
        past_date = today - timedelta(days=180)
    else:
        return None
        
    return past_date.strftime("%Y-%m-%dT%H:%M:%SZ")

# Function to get popular videos from YouTube
def get_popular_videos(api_key, query, max_results, published_after):
    try:
        youtube = build("youtube", "v3", developerKey=api_key)
        
        # First get video IDs from search
        search_request = youtube.search().list(
            part="id,snippet",
            q=query,
            type="video",
            order="viewCount",
            publishedAfter=published_after,
            maxResults=max_results
        )
        search_response = search_request.execute()
        
        # Extract video IDs
        video_ids = [item['id']['videoId'] for item in search_response['items']]
        
        # Get detailed video statistics
        videos_request = youtube.videos().list(
            part="snippet,statistics",
            id=','.join(video_ids)
        )
        videos_response = videos_request.execute()
        
        # Process and return the results
        results = []
        for item in videos_response['items']:
            results.append({
                'title': item['snippet']['title'],
                'channel': item['snippet']['channelTitle'],
                'published_at': item['snippet']['publishedAt'],
                'view_count': int(item['statistics'].get('viewCount', 0)),
                'like_count': int(item['statistics'].get('likeCount', 0)),
                'comment_count': int(item['statistics'].get('commentCount', 0)),
                'video_id': item['id'],
                'thumbnail': item['snippet']['thumbnails']['high']['url'],
                'description': item['snippet']['description']
            })
        
        # Sort by view count
        results.sort(key=lambda x: x['view_count'], reverse=True)
        return results
    
    except Exception as e:
        st.error(f"Error fetching YouTube data: {str(e)}")
        return []

# Function to generate brief context for each video
def generate_video_context(title, description):
    # Extract first 200 characters of description or less
    brief_desc = description[:200] + "..." if len(description) > 200 else description
    
    # Simple rule-based contextualizing (could be replaced with more sophisticated NLP)
    context = ""
    
    if re.search(r'meditation|mindfulness', title.lower() + brief_desc.lower()):
        context = "Meditation/Mindfulness practice"
    elif re.search(r'buddhis|zen|tao', title.lower() + brief_desc.lower()):
        context = "Eastern philosophy"
    elif re.search(r'christian|jesus|bible|faith', title.lower() + brief_desc.lower()):
        context = "Christian spirituality"
    elif re.search(r'islam|muslim|quran', title.lower() + brief_desc.lower()):
        context = "Islamic spirituality"
    elif re.search(r'judaism|jewish|torah', title.lower() + brief_desc.lower()):
        context = "Jewish spirituality"
    elif re.search(r'hindu|vedanta|yoga', title.lower() + brief_desc.lower()):
        context = "Hindu spirituality"
    elif re.search(r'consciousness|awareness', title.lower() + brief_desc.lower()):
        context = "Consciousness exploration"
    elif re.search(r'psychedelic|plant medicine|ayahuasca|dmt', title.lower() + brief_desc.lower()):
        context = "Psychedelic spirituality"
    elif re.search(r'near death|afterlife|heaven', title.lower() + brief_desc.lower()):
        context = "Afterlife exploration"
    elif re.search(r'science|physics|quantum', title.lower() + brief_desc.lower()):
        context = "Science and spirituality"
    elif re.search(r'gnosticism|gnostic|consciousness|awareness', title.lower() + brief_desc.lower()):
        context = "Gnosticism"
    else:
        context = "General spiritual content"
    
    return context

# NEW JSON-based function to generate lecture themes
def generate_lecture_themes_json(api_key, video_data, age_group):
    """
    Generate lecture themes using Gemini API with structured JSON output
    to avoid parsing issues later.
    
    Parameters:
    api_key (str): Gemini API key
    video_data (list): List of dicts containing video title and context
    age_group (str): Target age group (e.g., "20-30", "30-40", etc.)
    
    Returns:
    list: A list of theme dictionaries with structured data
    """
    try:
        # Configure the Gemini API
        genai.configure(api_key=api_key)
        
        # Initialize with the appropriate model
        try:
            models = genai.list_models()
            # Try to find gemini-2.0-flash or the best available model
            gemini_model = None
            for model_name in [m.name for m in models]:
                if 'gemini-2.0-flash' in model_name:
                    gemini_model = model_name
                    break
            if not gemini_model:
                for model_name in [m.name for m in models]:
                    if 'gemini' in model_name:
                        gemini_model = model_name
                        break
            if not gemini_model:
                gemini_model = 'gemini-2.0-flash'  # Default fallback
        except:
            gemini_model = 'gemini-2.0-flash'  # Default if we can't list models
        
        model = genai.GenerativeModel(gemini_model)
        
        # Prepare prompt with video titles and contexts
        titles_context = "\n".join([f"- {video['title']} ({video['context']})" for video in video_data])
        
        # Define age group characteristics
        age_characteristics = {
            "20-30": "digital natives, social media focused, seeking authenticity, concerned about climate crisis, mental health aware",
            "30-40": "career-focused, starting families, balancing work-life, health conscious, pragmatic spirituality",
            "40-50": "mid-life reflection, established careers, parenting teens, seeking deeper meaning, stress management",
            "50-60": "empty nest transitions, career peak or change, caring for aging parents, legacy considerations",
            "60+": "retirement planning/living, health challenges, grandparenting, mortality awareness, wisdom sharing"
        }
        
        # Get philosophical context if available
        philosophy_context = ""
        if 'philosophy_context_cleaned' in st.session_state and st.session_state['philosophy_context_cleaned']:
            full_context = st.session_state['philosophy_context_cleaned']
            # Limit context length to avoid token limits
            if len(full_context) > 10000:
                philosophy_context = full_context[:10000] + "..."
            else:
                philosophy_context = full_context
        
        # New prompt that requests JSON output
        if philosophy_context:
            prompt = f"""
As a spiritual content creator for a philosophical school of thought, analyze these trending YouTube video titles related to spirituality:

{titles_context}

The philosophical school has the following context, which should guide your suggestions:
----
{philosophy_context}
----

Based on these trends and the philosophical context, create 10 compelling lecture themes that would resonate specifically with people aged {age_group} years. 
Consider that this age group typically has these characteristics: {age_characteristics.get(age_group, "")}.

Make sure your suggested themes align with the philosophical approach described in the context.

IMPORTANT: Return your response in a valid JSON format with an array of 10 theme objects. Each theme object must have these exact fields:
{{
  "title": "The catchy title that reflects both current trends and philosophical approach",
  "description": "A short description (2-3 sentences)",
  "age_resonance": "Explanation of why this theme resonates with this specific age group",
  "philosophical_connection": "Brief note on how it connects to the philosophical context",
  "lecture_outline": "An outline for a 30-minute lecture based on the theme",
  "teaser": "A 50-60 word teaser that would be compelling for marketing purposes",
  "full_text": "A 500-word text that expands on the theme for a document/flyer"
}}

Make sure all fields are properly escaped for valid JSON and that the entire response is a valid JSON array.
"""
        else:
            # JSON prompt without philosophical context
            prompt = f"""
As a spiritual content creator, analyze these trending YouTube video titles related to spirituality:

{titles_context}

Based on these trends, create 10 compelling lecture themes that would resonate specifically with people aged {age_group} years. 
Consider that this age group typically has these characteristics: {age_characteristics.get(age_group, "")}.

IMPORTANT: Return your response in a valid JSON format with an array of 10 theme objects. Each theme object must have these exact fields:
{{
  "title": "The catchy title that reflects both current trends and philosophical approach",
  "description": "A short description (2-3 sentences)",
  "age_resonance": "Explanation of why this theme resonates with this specific age group",
  "philosophical_connection": "Brief note on how it connects to current spiritual trends",
  "lecture_outline": "An outline for a 30-minute lecture based on the theme",
  "teaser": "A 50-60 word teaser that would be compelling for marketing purposes",
  "full_text": "A 500-word text that expands on the theme for a document/flyer"
}}

Make sure all fields are properly escaped for valid JSON and that the entire response is a valid JSON array.
"""

        # Generate the response
        response = model.generate_content(prompt)
        
        # Extract the text from the response
        raw_response = response.text
        
        # Store the raw response for display
        result_text = raw_response
        
        # Sanitize and parse JSON response
        try:
            # First, find JSON array in the response (in case there's explanatory text)
            json_match = re.search(r'\[(.*?)\]', raw_response, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
            else:
                json_str = raw_response
            
            # Parse the JSON
            themes = json.loads(json_str)
            
            # Store in session state
            st.session_state['generated_themes'] = themes
            
            return themes, result_text
            
        except json.JSONDecodeError:
            # If JSON parsing fails, try to fix common issues
            st.warning("JSON parsing failed, attempting to fix formatting...")
            
            # Try to extract content between triple backticks if present
            code_block_match = re.search(r'```(?:json)?(.*?)```', raw_response, re.DOTALL)
            if code_block_match:
                json_str = code_block_match.group(1).strip()
            else:
                json_str = raw_response
            
            # Try manual extraction if still not valid JSON
            try:
                themes = json.loads(json_str)
            except:
                # Fall back to manual extraction if JSON parsing fails
                st.error("JSON parsing failed. Creating basic theme structure manually.")
                
                # Attempt to manually extract themes
                themes = []
                theme_matches = re.finditer(r'(?:"title"|{)\s*:\s*"([^"]+)"', json_str)
                
                for i, match in enumerate(theme_matches, 1):
                    if i > 10:  # Limit to 10 themes
                        break
                    
                    theme_title = match.group(1)
                    themes.append({
                        "title": theme_title,
                        "description": f"Description for '{theme_title}'",
                        "teaser": f"Teaser for '{theme_title}'",
                        "full_text": f"This is a placeholder for the full text about '{theme_title}'."
                    })
            
            # Store in session state
            st.session_state['generated_themes'] = themes
            return themes, result_text
    
    except Exception as e:
        st.error(f"Error generating lecture themes: {str(e)}")
        return [], str(e)

# Function to parse themes from text (keep for backward compatibility)
def parse_themes_from_text(themes_text):
    """
    Parse the generated themes text into a list of theme dictionaries.
    Each theme should have 'title', 'teaser', and 'full_text' fields.
    """
    themes = []
    
    # Split the text by theme sections (usually numbered themes)
    # This regex looks for patterns like "1. Theme Title" or "Theme 1: Title"
    theme_sections = re.split(r'\n\s*(?:\d+\.\s*|\bTheme\s+\d+[\s:]+)', themes_text)
    
    # Remove empty sections
    theme_sections = [section.strip() for section in theme_sections if section.strip()]
    
    for i, section in enumerate(theme_sections):
        theme = {}
        
        # Extract title
        title_match = re.search(r'^(.*?)(?:\n|$)', section)
        if title_match:
            theme['title'] = title_match.group(1).strip()
        else:
            theme['title'] = f"Theme {i+1}"
        
        # Extract teaser (look for text between asterisks or labeled as "Teaser:")
        teaser_match = re.search(r'\*\*(.*?)\*\*', section) or \
                      re.search(r'(?:Teaser|Short description):\s*(.*?)(?=\n\n|\n[A-Z])', section, re.DOTALL)
        if teaser_match:
            theme['teaser'] = teaser_match.group(1).strip()
        
        # Extract full text (everything after teaser or title if no teaser)
        if 'teaser' in theme:
            # Try to find text after the teaser
            full_text_match = re.search(re.escape(theme['teaser']) + r'\s*\n\n(.*)', section, re.DOTALL)
            if full_text_match:
                theme['full_text'] = full_text_match.group(1).strip()
        
        # If no full_text was found, try to extract the largest paragraph
        if 'full_text' not in theme or not theme['full_text']:
            paragraphs = re.split(r'\n\n+', section)
            # Get the longest paragraph that's not the title or teaser
            candidate_paragraphs = []
            for para in paragraphs:
                para = para.strip()
                if para and para != theme.get('title', '') and para != theme.get('teaser', ''):
                    candidate_paragraphs.append(para)
            
            if candidate_paragraphs:
                # Use the longest paragraph as full_text
                theme['full_text'] = max(candidate_paragraphs, key=len)
            else:
                # If no suitable paragraph, create a default
                theme['full_text'] = """This theme explores the connection between modern technology and ancient wisdom, 
                offering insights on how to navigate today's digital world while maintaining spiritual authenticity. 
                Drawing from timeless principles, we'll discuss practical techniques for personal growth, 
                mindfulness, and deeper connection in an age of distraction."""
        
        themes.append(theme)
    
    # If no themes were parsed, create a default one
    if not themes:
        themes = [{
            'title': 'The Algorithm of Your Soul: Breaking Free from the Echo Chamber',
            'teaser': 'Feeling trapped in an online echo chamber? Discover the \'algorithm of your soul\' – the invisible code dictating your thoughts and actions.',
            'full_text': """The algorithm of your soul is the invisible pattern that shapes your thoughts, behaviors, and perceptions. Just as digital algorithms influence what you see online, internal algorithms determine how you experience life.

By understanding these inner patterns, you can rewrite the code that limits your potential. Ancient wisdom traditions have long recognized these patterns and offer timeless techniques to transcend them.

This journey of awakening begins with awareness - recognizing when you're trapped in echo chambers of your own making. Then, through mindful reflection and intentional practice, you can gradually reprogram these patterns for a more authentic, fulfilling life.

Break free from the digital matrix. Discover your true potential."""
        }]
    
    return themes

# Older function for legacy support
def generate_lecture_themes(api_key, video_data, age_group):
    # Call the JSON version and convert result
    themes, raw_text = generate_lecture_themes_json(api_key, video_data, age_group)
    return raw_text


def translate_themes_to_portuguese(api_key, themes):
    """
    Translate the generated themes from English to Portuguese using Gemini API.
    
    Parameters:
    api_key (str): Gemini API key
    themes (list): List of theme dictionaries in English
    
    Returns:
    tuple: (portuguese_themes, raw_json_text)
    """
    try:
        # Configure the Gemini API
        genai.configure(api_key=api_key)
        
        # Initialize with the appropriate model
        try:
            models = genai.list_models()
            # Try to find a suitable model for translation
            gemini_model = None
            for model_name in [m.name for m in models]:
                if 'gemini-2.0-flash' in model_name:
                    gemini_model = model_name
                    break
            if not gemini_model:
                for model_name in [m.name for m in models]:
                    if 'gemini' in model_name:
                        gemini_model = model_name
                        break
            if not gemini_model:
                gemini_model = 'gemini-2.0-flash'  # Default fallback
        except:
            gemini_model = 'gemini-2.0-flash'  # Default if we can't list models
        
        model = genai.GenerativeModel(gemini_model)
        
        # Convert the themes to JSON for the translation prompt
        english_json = json.dumps(themes, ensure_ascii=False, indent=2)
        
        # Create a prompt for translation
        prompt = f"""
You are a professional translator with expertise in spirituality, philosophy, and psychology.

Translate the following JSON data containing lecture themes from English to Portuguese. Maintain the exact same JSON structure and field names, but translate all content values.

Pay special attention to properly translating spiritual and philosophical terms. Ensure the Portuguese translation maintains the spiritual essence and nuance of the original.

The translation should sound natural and idiomatic in Portuguese while preserving the meaning.

JSON to translate:
```json
{english_json}
```

IMPORTANT: Return ONLY the translated JSON array with NO additional text or explanation. The result must be valid JSON that can be parsed programmatically.
"""

        # Generate the translation
        response = model.generate_content(prompt)
        
        # Extract the text from the response
        raw_response = response.text
        
        # Try to extract JSON from the response
        try:
            # First, find JSON array in the response (in case there's explanatory text)
            json_match = re.search(r'\[(.*?)\]', raw_response, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
            else:
                json_str = raw_response
            
            # Look for JSON content between code blocks
            code_block_match = re.search(r'```(?:json)?(.*?)```', raw_response, re.DOTALL)
            if code_block_match:
                json_str = code_block_match.group(1).strip()
            
            # Parse the JSON
            portuguese_themes = json.loads(json_str)
            
            return portuguese_themes, raw_response
            
        except json.JSONDecodeError:
            # If JSON parsing fails, log error and return the original themes
            print(f"Failed to parse translated JSON: {raw_response[:500]}...")
            return themes, raw_response
    
    except Exception as e:
        print(f"Error translating themes: {str(e)}")
        return themes, str(e)









def create_theme_document_with_language_option(theme, gemini_api_key, language="english"):
    """
    Creates a formatted Word document with the theme content in the selected language.
    
    Parameters:
    theme (dict): Dictionary containing theme data with consistent keys
    gemini_api_key (str): API key for Gemini used for image generation
    language (str): "english" or "portuguese"
    
    Returns:
    BytesIO: A BytesIO object containing the generated Word document
    """
    # Create a new Word document
    doc = docx.Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Generate a query based on the theme title and add some spiritual keywords
    try:
        # Extract meaningful information from the theme
        clean_title = theme.get('title', '').lower()
        description = theme.get('description', '').lower()
        
        # Extract key concepts
        key_words = []
        
        # Extract nouns and adjectives from title
        title_words = clean_title.replace('-', ' ').split()
        # Filter out common stop words
        stop_words = ['the', 'and', 'of', 'in', 'to', 'a', 'is', 'that', 'it', 'with', 'as', 'for',
                     'o', 'a', 'os', 'as', 'de', 'da', 'do', 'das', 'dos', 'em', 'no', 'na', 'um', 'uma']
        key_words.extend([word for word in title_words if word not in stop_words and len(word) > 3])
        
        # Create a list of spiritual keywords for better image results
        spiritual_keywords = [
            "spiritual", "meditation", "abstract", "enlightenment", "mindfulness", "consciousness",
            "sacred", "divine", "cosmic", "transcendence", "wisdom", "harmony", "balance",
            "serenity", "energy", "light", "nature", "universe", "soul", "spirit"
        ]
        
        # Select 2-3 random keywords to add variety
        import random
        selected_keywords = random.sample(spiritual_keywords, min(3, len(spiritual_keywords)))
        
        # If we have key words from the title, prioritize those and add just 1-2 spiritual keywords
        if key_words:
            # Combine up to 2 key words from title with 1-2 spiritual keywords
            image_query = " ".join(key_words[:2] + selected_keywords[:2])
        else:
            # If no good key words from title, use more spiritual keywords
            image_query = " ".join(selected_keywords[:3])
        
        # Add "abstract art" to get more artistic images
        image_query += " abstract art"
        
        # Format for Unsplash API
        formatted_query = image_query.replace(' ', '+')
        
        # Use a more targeted Unsplash query
        image_url = f"https://source.unsplash.com/1200x600/?{formatted_query}"
        
        # Get the image
        img_response = requests.get(image_url)
        img = Image.open(BytesIO(img_response.content))
        
        # Save the image to a temporary file
        img_path = "temp_theme_image.jpg"
        img.save(img_path)
        
        # Add the image to the document
        doc.add_picture(img_path, width=Inches(6))
        os.remove(img_path)  # Clean up temp file
        
        # Add some space after the image
        doc.add_paragraph()
    except Exception as e:
        # If image acquisition fails, log and continue without an image
        print(f"Image generation failed: {str(e)}")
    
    # Add the theme title
    title_paragraph = doc.add_paragraph()
    clean_title = theme.get('title', 'Theme Title' if language == "english" else "Título do Tema")
    title_run = title_paragraph.add_run(clean_title)
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Get month names in Portuguese if needed
    pt_months = {
        "January": "Janeiro", "February": "Fevereiro", "March": "Março", 
        "April": "Abril", "May": "Maio", "June": "Junho",
        "July": "Julho", "August": "Agosto", "September": "Setembro",
        "October": "Outubro", "November": "Novembro", "December": "Dezembro"
    }
    
    # Add the date (current month and year)
    date_paragraph = doc.add_paragraph()
    date_str = datetime.now().strftime("%B %Y")
    
    # Translate month if using Portuguese
    if language == "portuguese":
        month_name = datetime.now().strftime("%B")
        if month_name in pt_months:
            date_str = date_str.replace(month_name, pt_months[month_name])
    
    date_run = date_paragraph.add_run(date_str.upper())
    date_run.bold = True
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add the teaser
    teaser_content = theme.get('teaser', '')
    if teaser_content:
        doc.add_paragraph()
        teaser_paragraph = doc.add_paragraph()
        teaser_run = teaser_paragraph.add_run(teaser_content)
        teaser_run.italic = True
        teaser_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add the full text
    full_text = theme.get('full_text', '')
    if not full_text:
        # Default text if no full_text is found
        if language == "english":
            full_text = """The algorithm of your soul is the invisible pattern that shapes your thoughts, behaviors, and perceptions. Just as digital algorithms influence what you see online, internal algorithms determine how you experience life."""
        else:
            full_text = """O algoritmo da sua alma é o padrão invisível que molda seus pensamentos, comportamentos e percepções. Assim como os algoritmos digitais influenciam o que você vê online, os algoritmos internos determinam como você experimenta a vida."""
    
    # Add the main text
    doc.add_paragraph()  # Add space
    
    # Split into paragraphs and add them
    paragraphs = re.split(r'\n\s*\n', full_text)
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
            
        # Skip any lines that look like outline points
        if re.match(r'^[-•*]\s', paragraph) or re.match(r'^\d+\.\s', paragraph):
            continue
            
        # Add as regular paragraph
        p = doc.add_paragraph()
        p.add_run(paragraph.strip())
    
    # Add footer with Rosacruz Áurea branding
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_run = footer_paragraph.add_run("ROSACRUZ ÁUREA | LECTORIUM ROSICRUCIANUM")
    footer_run.font.size = Pt(8)
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document to a BytesIO object
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

# Keep original function for backward compatibility
def create_theme_document(theme):
    """
    Creates a formatted Word document based on a theme dictionary.
    Now uses the more robust create_theme_document_from_json
    """
    return create_theme_document_from_json(theme)

# Function to download a document
def get_binary_file_downloader_html(bin_data, file_label='Document', file_name='theme.docx'):
    b64_data = base64.b64encode(bin_data.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_data}" download="{file_name}">{file_label}</a>'

# Main app layout
tab1, tab2, tab3 = st.tabs(["Mine YouTube Videos", "Lecture Theme Generator", "About"])

with tab1:
    st.header("Mine Popular Spirituality Videos")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.subheader("Last Week")
        if st.button("Mine Last Week's Videos"):
            if youtube_api_key:
                with st.spinner("Fetching last week's popular videos..."):
                    published_after = get_date_for_period("1 week")
                    videos = get_popular_videos(youtube_api_key, search_query, max_results, published_after)
                    
                    if videos:
                        # Add context to each video
                        for video in videos:
                            video['context'] = generate_video_context(video['title'], video['description'])
                        
                        # Store in session state for later use
                        st.session_state['weekly_videos'] = videos
                        
                        # Display videos
                        for i, video in enumerate(videos):
                            st.write(f"**{i+1}. {video['title']}**")
                            st.write(f"*Context: {video['context']}*")
                            st.write(f"Views: {video['view_count']:,} | Channel: {video['channel']}")
                            st.write(f"[Watch on YouTube](https://www.youtube.com/watch?v={video['video_id']})")
                            st.image(video['thumbnail'], use_container_width=True)
                            st.divider()
                    else:
                        st.warning("No videos found or error occurred.")
            else:
                st.error("Please enter your YouTube API key in the sidebar.")
    
    with col2:
        st.subheader("Last Month")
        if st.button("Mine Last Month's Videos"):
            if youtube_api_key:
                with st.spinner("Fetching last month's popular videos..."):
                    published_after = get_date_for_period("1 month")
                    videos = get_popular_videos(youtube_api_key, search_query, max_results, published_after)
                    
                    if videos:
                        # Add context to each video
                        for video in videos:
                            video['context'] = generate_video_context(video['title'], video['description'])
                        
                        # Store in session state for later use
                        st.session_state['monthly_videos'] = videos
                        
                        # Display videos
                        for i, video in enumerate(videos):
                            st.write(f"**{i+1}. {video['title']}**")
                            st.write(f"*Context: {video['context']}*")
                            st.write(f"Views: {video['view_count']:,} | Channel: {video['channel']}")
                            st.write(f"[Watch on YouTube](https://www.youtube.com/watch?v={video['video_id']})")
                            st.image(video['thumbnail'], use_container_width=True)
                            st.divider()
                    else:
                        st.warning("No videos found or error occurred.")
            else:
                st.error("Please enter your YouTube API key in the sidebar.")
    
    with col3:
        st.subheader("Last 6 Months")
        if st.button("Mine Last 6 Months' Videos"):
            if youtube_api_key:
                with st.spinner("Fetching last 6 months' popular videos..."):
                    published_after = get_date_for_period("6 months")
                    videos = get_popular_videos(youtube_api_key, search_query, max_results, published_after)
                    
                    if videos:
                        # Add context to each video
                        for video in videos:
                            video['context'] = generate_video_context(video['title'], video['description'])
                        
                        # Store in session state for later use
                        st.session_state['biannual_videos'] = videos
                        
                        # Display videos
                        for i, video in enumerate(videos):
                            st.write(f"**{i+1}. {video['title']}**")
                            st.write(f"*Context: {video['context']}*")
                            st.write(f"Views: {video['view_count']:,} | Channel: {video['channel']}")
                            st.write(f"[Watch on YouTube](https://www.youtube.com/watch?v={video['video_id']})")
                            st.image(video['thumbnail'], use_container_width=True)
                            st.divider()
                    else:
                        st.warning("No videos found or error occurred.")
            else:
                st.error("Please enter your YouTube API key in the sidebar.")

with tab2:
    st.header("Generate Lecture Themes by Age Group")
    
    # Show philosophy context status
    if 'philosophy_context_cleaned' in st.session_state and st.session_state['philosophy_context_cleaned']:
        context_length = len(st.session_state['philosophy_context_cleaned'])
        st.success(f"✅ Philosophy context loaded ({context_length} characters)")
        with st.expander("View loaded philosophical context"):
            st.write(st.session_state['philosophy_context_cleaned'][:2700] + "..." 
                    if context_length > 2700 else st.session_state['philosophy_context_cleaned'])
    else:
        st.warning("⚠️ No philosophical context loaded. Upload an HTML file in the sidebar to provide context.")
    
    # Select data source
    data_source = st.selectbox(
        "Select Video Data Source",
        ["Last Week", "Last Month", "Last 6 Months", "Combined (All Time Periods)"]
    )
    
    # Map selection to session state keys
    data_mapping = {
        "Last Week": 'weekly_videos',
        "Last Month": 'monthly_videos',
        "Last 6 Months": 'biannual_videos'
    }
    
    # Get videos from selected source
    selected_videos = []
    if data_source in data_mapping:
        selected_videos = st.session_state.get(data_mapping[data_source], [])
    elif data_source == "Combined (All Time Periods)":
        # Combine all video sources
        weekly = st.session_state.get('weekly_videos', [])
        monthly = st.session_state.get('monthly_videos', [])
        biannual = st.session_state.get('biannual_videos', [])
        
        # Add source field to track which time period each video came from
        for v in weekly:
            v['source'] = 'Last Week'
        for v in monthly:
            v['source'] = 'Last Month'
        for v in biannual:
            v['source'] = 'Last 6 Months'
        
        # Combine and deduplicate by video ID
        all_videos = weekly + monthly + biannual
        video_ids_seen = set()
        selected_videos = []
        
        for video in all_videos:
            if video['video_id'] not in video_ids_seen:
                selected_videos.append(video)
                video_ids_seen.add(video['video_id'])
    
    # Show summary of available videos
    if selected_videos:
        st.write(f"Found {len(selected_videos)} videos from {data_source}")
        
        # Display a sample of video titles (first 10)
        st.write("Sample of video titles:")
        for i, video in enumerate(selected_videos[:10]):
            st.write(f"{i+1}. {video['title']} - *{video['context']}*")
        
        if len(selected_videos) > 10:
            st.write(f"...and {len(selected_videos)-10} more")
    else:
        st.warning(f"No videos available for {data_source}. Please mine videos in the first tab.")
    
    # Age group selection
    age_group = st.selectbox(
        "Select Target Age Group",
        ["20-30", "30-40", "40-50", "50-60", "60+"]
    )
    
    # Generate themes button
    if st.button("Generate Lecture Themes"):
        if gemini_api_key and selected_videos:
            with st.spinner(f"Generating lecture themes for {age_group} age group..."):
                try:
                    # Use the new JSON-based theme generator
                    themes, themes_text = generate_lecture_themes_json(gemini_api_key, selected_videos, age_group)
                    
                    # Check if themes exists and is not empty
                    if themes:
                        st.markdown("## Generated Themes")
                        
                        # Display the raw JSON in an expander for debugging
                        #with st.expander("View Raw JSON Response"):
                        #    st.code(themes_text)

                        # Translate themes to Portuguese
                        with st.spinner("Translating themes to Portuguese..."):
                            portuguese_themes, portuguese_text = translate_themes_to_portuguese(gemini_api_key, themes)
                            
                            # Store Portuguese themes in session state
                            st.session_state['portuguese_themes'] = portuguese_themes
                            
                            # Display the Portuguese JSON
                            #with st.expander("View Raw JSON Response (Portuguese)"):
                            #    st.code(portuguese_text)
  



                        
                        # Display the parsed themes
                        tabs = st.tabs(["English", "Portuguese"])
                        
                        with tabs[0]:  # English tab
                            
                            for i, theme in enumerate(themes):
                                st.markdown(f"### {i+1}. {theme.get('title', 'Untitled Theme')}")
                                
                                if 'description' in theme:
                                    st.markdown(f"**Description:** {theme['description']}")
                                
                                if 'teaser' in theme:
                                    st.markdown(f"**Teaser:** *{theme['teaser']}*")
                                
                                with st.expander("View Full Details"):
                                    if 'age_resonance' in theme:
                                        st.markdown(f"**Age Group Resonance:** {theme['age_resonance']}")
                                    
                                    if 'philosophical_connection' in theme:
                                        st.markdown(f"**Philosophical Connection:** {theme['philosophical_connection']}")
                                    
                                    if 'lecture_outline' in theme:
                                        st.markdown(f"**Lecture Outline:**\n{theme['lecture_outline']}")
                                    
                                    if 'full_text' in theme:
                                        st.markdown(f"**Full Text:**\n{theme['full_text']}")

                        with tabs[1]:  # Portuguese tab
                            for i, theme in enumerate(portuguese_themes):
                                st.markdown(f"### {i+1}. {theme.get('title', 'Tema Sem Título')}")
                                
                                if 'description' in theme:
                                    st.markdown(f"**Descrição:** {theme['description']}")
                                
                                if 'teaser' in theme:
                                    st.markdown(f"**Chamada:** *{theme['teaser']}*")
                                
                                with st.expander("Ver Detalhes Completos"):
                                    if 'age_resonance' in theme:
                                        st.markdown(f"**Ressonância com a Faixa Etária:** {theme['age_resonance']}")
                                    
                                    if 'philosophical_connection' in theme:
                                        st.markdown(f"**Conexão Filosófica:** {theme['philosophical_connection']}")
                                    
                                    if 'lecture_outline' in theme:
                                        st.markdown(f"**Estrutura da Palestra:**\n{theme['lecture_outline']}")
                                    
                                    if 'full_text' in theme:
                                        st.markdown(f"**Texto Completo:**\n{theme['full_text']}")
                        
                except Exception as e:
                    st.error(f"Error generating lecture themes: {str(e)}")
                    st.exception(e)  # This will show the full traceback
                    
                    # Create a default theme even on error
                    st.session_state['generated_themes'] = [{
                        'title': 'Default Theme (Error Recovery)',
                        'teaser': 'A placeholder theme created when an error occurred.',
                        'full_text': """This is a default theme created when an error occurred during theme generation. You can still use this to test document generation."""
                    }]
        elif not gemini_api_key:
            st.error("Please ensure your Google Gemini API key is properly set.")
        else:
            st.error("No video data available. Please mine videos first.")
                
    # Show theme details and document generation if we have generated themes
    if 'generated_themes' in st.session_state and st.session_state['generated_themes']:
        st.markdown("---")
        st.markdown("## Create Document for Theme")
        
        # Check if we have Portuguese themes
        has_portuguese = 'portuguese_themes' in st.session_state and st.session_state['portuguese_themes']
        
        # Language selection option
        language_option = "english"
        if has_portuguese:
            language_option = st.radio(
                "Select Document Language",
                ["English", "Portuguese"],
                horizontal=True
            )
            language_option = language_option.lower()
        
        # Determine which themes to use based on language
        if language_option == "portuguese" and has_portuguese:
            themes_to_use = st.session_state['portuguese_themes']
        else:
            themes_to_use = st.session_state['generated_themes']
        
        # If we have valid themes, proceed with selection and display
        if themes_to_use and len(themes_to_use) > 0:
            # Create options for selectbox - extract just the titles
            options = []
            for i, theme in enumerate(themes_to_use):
                # Get title or fallback to a default
                if isinstance(theme, dict) and 'title' in theme:
                    # Clean up the title if needed
                    title = theme['title']
                    # Remove any markdown formatting
                    title = title.replace('*', '').replace('#', '').strip()
                    options.append(title)
                else:
                    options.append(f"Theme {i+1}" if language_option == "english" else f"Tema {i+1}")
            











            # Replace the single selectbox with multiselect for multiple theme selection
            select_label = "Select themes to create documents" if language_option == "english" else "Selecione temas para criar documentos"
            selected_options = st.multiselect(
                select_label,
                options,
                max_selections=10  # Allow up to 10 selections
            )

            # Check if any themes are selected
            if selected_options:
                # Find all selected themes
                selected_themes = []
                for selected_option in selected_options:
                    for theme in themes_to_use:
                        if isinstance(theme, dict) and 'title' in theme:
                            clean_title = theme['title'].replace('*', '').replace('#', '').strip()
                            if clean_title == selected_option:
                                selected_themes.append((selected_option, theme))
                                break
                
                # Create buttons for individual and batch downloads
                col1, col2 = st.columns(2)
                
                with col1:
                    # Button for generating and downloading individual documents
                    button_label = "Generate Individual Documents" if language_option == "english" else "Gerar Documentos Individuais"
                    individual_btn = st.button(button_label)
                
                with col2:
                    # Button for generating and downloading all documents in a zip
                    zip_button_label = "Generate & Download All as ZIP" if language_option == "english" else "Gerar & Baixar Todos como ZIP"
                    zip_btn = st.button(zip_button_label)
                
                # Handle individual document generation
                if individual_btn:
                    spinner_text = "Creating documents with automatically generated images..." if language_option == "english" else "Criando documentos com imagens geradas automaticamente..."
                    
                    with st.spinner(spinner_text):
                        # Create a container for all download buttons
                        download_container = st.container()
                        success_count = 0
                        
                        # Process each selected theme
                        for selected_option, selected_theme in selected_themes:
                            try:
                                # Use the document creator function with language option
                                doc_bytes = create_theme_document_with_language_option(
                                    selected_theme, 
                                    gemini_api_key,
                                    language_option
                                )
                                
                                # Sanitize filename
                                safe_title = re.sub(r'[^\w\-_\. ]', '', selected_option)
                                safe_title = safe_title.replace(' ', '_')
                                lang_suffix = "" if language_option == "english" else "_pt"
                                filename = f"{safe_title}{lang_suffix}_{datetime.now().strftime('%Y%m%d')}.docx"
                                
                                # Add download button for this document
                                with download_container:
                                    download_label = f"Download: {selected_option}" if language_option == "english" else f"Baixar: {selected_option}"
                                    st.download_button(
                                        label=download_label,
                                        data=doc_bytes,
                                        file_name=filename,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key=f"download_{safe_title}"  # Unique key for each button
                                    )
                                
                                success_count += 1
                                
                            except Exception as e:
                                error_msg = f"Error creating document for '{selected_option}': {str(e)}" if language_option == "english" else f"Erro ao criar documento para '{selected_option}': {str(e)}"
                                st.error(error_msg)
                        
                        # Show a summary message
                        if success_count > 0:
                            success_msg = f"{success_count} documents created successfully!" if language_option == "english" else f"{success_count} documentos criados com sucesso!"
                            st.success(success_msg)
                
                # Handle zip file generation
                elif zip_btn:
                    zip_spinner_text = "Creating all documents and preparing ZIP file..." if language_option == "english" else "Criando todos os documentos e preparando arquivo ZIP..."
                    
                    with st.spinner(zip_spinner_text):
                        # Create a BytesIO object to store the zip file
                        zip_buffer = io.BytesIO()
                        
                        # Create a ZipFile object
                        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                            success_count = 0
                            error_docs = []
                            
                            # Process each selected theme
                            for selected_option, selected_theme in selected_themes:
                                try:
                                    # Use the document creator function with language option
                                    doc_bytes = create_theme_document_with_language_option(
                                        selected_theme, 
                                        gemini_api_key,
                                        language_option
                                    )
                                    
                                    # Sanitize filename
                                    safe_title = re.sub(r'[^\w\-_\. ]', '', selected_option)
                                    safe_title = safe_title.replace(' ', '_')
                                    lang_suffix = "" if language_option == "english" else "_pt"
                                    filename = f"{safe_title}{lang_suffix}_{datetime.now().strftime('%Y%m%d')}.docx"
                                    
                                    # Add file to zip
                                    zip_file.writestr(filename, doc_bytes.getvalue())
                                    success_count += 1
                                    
                                except Exception as e:
                                    error_docs.append(selected_option)
                        
                        # Set the buffer position to the beginning
                        zip_buffer.seek(0)
                        
                        # Show errors if any
                        if error_docs:
                            error_msg = f"Failed to generate {len(error_docs)} documents" if language_option == "english" else f"Falha ao gerar {len(error_docs)} documentos"
                            st.error(error_msg)
                        
                        # Show success message and download button if any documents were created
                        if success_count > 0:
                            success_msg = f"{success_count} documents created successfully!" if language_option == "english" else f"{success_count} documentos criados com sucesso!"
                            st.success(success_msg)
                            
                            # Create a download button for the zip file
                            date_str = datetime.now().strftime("%Y%m%d")
                            zip_label = "Download All Documents (ZIP)" if language_option == "english" else "Baixar Todos os Documentos (ZIP)"
                            zip_filename = f"lecture_themes_{language_option}_{date_str}.zip"
                            
                            st.download_button(
                                label=zip_label,
                                data=zip_buffer,
                                file_name=zip_filename,
                                mime="application/zip"
                            )

            # If no themes are selected, show message
            elif themes_to_use and len(themes_to_use) > 0:
                # Display a message to prompt selection
                prompt_msg = "Select themes above to create documents" if language_option == "english" else "Selecione temas acima para criar documentos"
                st.info(prompt_msg)
    




with tab3:
    st.header("About This Application")
    st.markdown("""
    ### Purpose
    This application is designed to help philosophical and spiritual educators identify trending topics 
    and create age-appropriate lecture content based on current interests in spirituality.
    
    ### How It Works
    1. **Data Mining**: The app connects to YouTube's API to find the most viewed videos on spirituality 
       from three different time periods.
    
    2. **Content Analysis**: Each video is briefly contextualized to identify its spiritual domain or approach.
    
    3. **Theme Generation**: Google's Gemini AI analyzes the trending content and suggests lecture themes 
       tailored to different generational perspectives and needs.
       
    4. **Document Creation**: For each generated theme, you can create a professional document with:
       - A thematic image
       - The theme title
       - A compelling 50-60 word teaser
       - A comprehensive 500-word text explaining the theme
    
    ### Requirements
    - YouTube Data API v3 key (get one from [Google Cloud Console](https://console.cloud.google.com/))
    - Google Gemini API key
    - Required Python packages: streamlit, pandas, google-api-python-client, google-generativeai, 
      python-docx, Pillow, BeautifulSoup4, requests
    
    ### Privacy
    This application does not store any data outside your browser session. API keys are not saved 
    between sessions for security reasons. Images for documents are fetched from Unsplash using search
    terms based on the theme title.
    """)

# Footer
st.divider()
st.caption("© 2025 Spirituality Trends Analyzer | Developed for philosophical education")